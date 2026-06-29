"""
AttendTrack — Documentation Generator (.docx)
Converts all markdown docs in docs/ into a single professional Word document.
"""

import os, re, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

DOCS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "AttendTrack-Dokumentasi-Sistem.docx")

# ── Colour Palette ────────────────────────────────────────────
PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)       # Deep navy blue
ACCENT = RGBColor(0x2C, 0x7B, 0xBE)         # Medium blue
LIGHT_BG = RGBColor(0xEB, 0xF2, 0xFA)       # Light blue bg
DARK_BG = RGBColor(0x1B, 0x3A, 0x5C)        # Dark bg
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
INFO_BG = RGBColor(0xE3, 0xF2, 0xFD)
WARN_BG = RGBColor(0xFF, 0xF3, 0xE0)
TIP_BG = RGBColor(0xE8, 0xF5, 0xE9)
SUCCESS_BG = RGBColor(0xE8, 0xF5, 0xE9)

doc = Document()

# ── Style Helpers ─────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [
    ('Heading 1', 22, PRIMARY), ('Heading 2', 16, ACCENT), ('Heading 3', 13, PRIMARY),
    ('Heading 4', 11, ACCENT),
]:
    s = doc.styles[level]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(18 if level == 'Heading 1' else 12)
    s.paragraph_format.space_after = Pt(8)
    if level == 'Heading 1':
        s.paragraph_format.page_break_before = False

# ── Helper functions ──────────────────────────────────────────
def add_horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="2C7BBE"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)

def add_page_break():
    doc.add_page_break()

def add_callout_box(text, box_type="info", title=None):
    """info / warning / tip"""
    bg = {'info': INFO_BG, 'warning': WARN_BG, 'tip': TIP_BG}[box_type]
    icon = {'info': 'INFO', 'warning': 'PENTING', 'tip': 'TIP'}[box_type]
    label = title or icon
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    # shading
    hex_color = str(bg)
    shading = parse_xml(
        '<w:shd %s w:fill="%s"/>' % (nsdecls('w'), hex_color)
    )
    cell._tc.get_or_add_tcPr().append(shading)
    # border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '<w:top w:val="single" w:sz="12" w:space="0" w:color="%s"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="%s"/>'
        '<w:right w:val="single" w:sz="12" w:space="0" w:color="%s"/>'
        '</w:tcBorders>' % (nsdecls('w'), ACCENT, ACCENT, ACCENT, ACCENT)
    )
    tcPr.append(tcBorders)
    # label
    p_label = cell.paragraphs[0]
    run = p_label.add_run(f"  {label}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    # content
    if '\n' in text:
        for line in text.split('\n'):
            if line.strip():
                p = cell.add_paragraph()
                run = p.add_run(line.strip())
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK
    else:
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
    doc.add_paragraph()  # spacing

def set_cell_shading(cell, color):
    shading = parse_xml(
        '<w:shd %s w:fill="%s"/>' % (nsdecls('w'), color)
    )
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, '1B3A5C')
    # rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if ri % 2 == 1:
                set_cell_shading(cell, 'EBF2FA')
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()

def add_code(text, language=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(
        '<w:shd %s w:fill="F5F5F5"/>' % nsdecls('w')
    )
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)

def add_numeric(text, level=0):
    p = doc.add_paragraph(text, style='List Number')

def add_summary_box(items):
    """Add a chapter summary box with check-mark items."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'EBF2FA')
    p = cell.paragraphs[0]
    run = p.add_run("  Ringkasan")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PRIMARY
    for item in items:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(f"  {item}")
        run2.font.size = Pt(10)
    doc.add_paragraph()

def add_icon_heading(icon, text, level=1):
    if level == 1:
        doc.add_heading(f"{icon} {text}", level=1)
    elif level == 2:
        doc.add_heading(f"{icon} {text}", level=2)
    else:
        doc.add_heading(f"{icon} {text}", level=3)

# ── COVER PAGE ────────────────────────────────────────────────
def build_cover():
    for _ in range(6):
        doc.add_paragraph()
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(14)
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ATTENDTRACK')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = PRIMARY
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sistem Absensi RFID + Face Recognition')
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(14)
    for _ in range(3):
        doc.add_paragraph()
    # Desc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Dokumentasi Sistem — Versi 3.0')
    run.font.size = Pt(13)
    run.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    add_page_break()

# ── KATA PENGANTAR ────────────────────────────────────────────
def build_kata_pengantar():
    doc.add_heading('Kata Pengantar', level=1)
    doc.add_paragraph(
        'Puji syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas terselesaikannya '
        'dokumentasi sistem AttendTrack — Sistem Absensi RFID + Face Recognition versi 3.0.'
    )
    doc.add_paragraph(
        'Dokumen ini disusun sebagai panduan lengkap yang mencakup seluruh aspek sistem, '
        'mulai dari arsitektur, komponen perangkat keras dan lunak, hingga panduan '
        'implementasi dan pemeliharaan. Dokumentasi ini ditujukan bagi berbagai kalangan, '
        'baik yang memiliki latar belakang teknis maupun non-teknis, termasuk dosen, '
        'manajer, stakeholder, klien, investor, auditor, dan pengguna umum.'
    )
    doc.add_paragraph(
        'AttendTrack menggabungkan teknologi RFID (Radio Frequency Identification) dengan '
        'kecerdasan buatan pengenalan wajah (Face Recognition AI) untuk memberikan solusi '
        'absensi dua faktor yang andal, akurat, dan real-time. Sistem ini dirancang untuk '
        'institusi pendidikan, perkantoran, dan lingkungan manufaktur yang membutuhkan '
        'pencatatan kehadiran yang terpercaya.'
    )
    doc.add_paragraph(
        'Kami berharap dokumentasi ini dapat memberikan pemahaman yang komprehensif dan '
        'menjadi referensi yang bermanfaat bagi semua pihak yang berkepentingan.'
    )
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Tim Pengembang AttendTrack')
    run.bold = True
    run.font.color.rgb = PRIMARY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(datetime.now().strftime('%B %Y'))
    run.font.color.rgb = GREY
    add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────
def build_toc():
    doc.add_heading('Daftar Isi', level=1)
    doc.add_paragraph('(Daftar isi otomatis akan muncul setelah dokumen dibuka di Microsoft Word — Perbarui dengan Ctrl+A lalu F9)')
    add_callout_box(
        'Untuk memperbarui daftar isi, klik kanan pada area daftar isi dan pilih "Update Field" '
        'atau tekan Ctrl+A lalu F9.',
        'tip'
    )
    # Manual TOC
    toc_items = [
        ('Kata Pengantar', ''),
        ('Executive Summary', ''),
        ('Bab 1: Gambaran Umum Proyek', '1'),
        ('Bab 2: Arsitektur Sistem', '2'),
        ('Bab 3: Backend', '3'),
        ('Bab 4: Frontend', '4'),
        ('Bab 5: Face Recognition', '5'),
        ('Bab 6: ESP8266 (RFID Scanner)', '6'),
        ('Bab 7: ESP32-CAM (Face Capture)', '7'),
        ('Bab 8: Database', '8'),
        ('Bab 9: API', '9'),
        ('Bab 10: Socket.IO', '10'),
        ('Bab 11: gRPC', '11'),
        ('Bab 12: Alur Aplikasi', '12'),
        ('Bab 13: Alur Fitur', '13'),
        ('Bab 14: Diagram Sekuens', '14'),
        ('Bab 15: Deployment', '15'),
        ('Bab 16: Troubleshooting', '16'),
        ('Bab 17: Praktik Terbaik', '17'),
        ('Kesimpulan', ''),
        ('Lampiran', ''),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, t in enumerate(['Bagian', 'Halaman']):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(t)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_cell_shading(hdr[i], '1B3A5C')
    for title, page in toc_items:
        row = table.add_row().cells
        row[0].text = ''
        run = row[0].paragraphs[0].add_run(title)
        run.font.size = Pt(10)
        row[1].text = ''
    add_page_break()

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────
def build_executive_summary():
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'AttendTrack adalah sistem manajemen absensi generasi terbaru yang mengintegrasikan '
        'teknologi RFID (Radio Frequency Identification) dengan kecerdasan buatan pengenalan '
        'wajah (Face Recognition AI). Sistem ini memberikan solusi absensi dua faktor yang '
        'andal, akurat, dan berjalan secara real-time.'
    )
    doc.add_paragraph(
        'Dikembangkan dengan arsitektur microservices, AttendTrack terdiri dari tiga komponen '
        'utama: (1) Express Gateway sebagai API Gateway dan REST API, (2) gRPC Server sebagai '
        'lapisan logika bisnis, dan (3) Python Face Recognition Service untuk pemrosesan '
        'pengenalan wajah berbasis deep learning.'
    )
    doc.add_paragraph(
        'Sistem ini didukung oleh perangkat IoT khusus — ESP8266 sebagai pemindai RFID '
        'dan ESP32-CAM sebagai kamera pengambil wajah — yang berkomunikasi secara nirkabel '
        'dengan server backend melalui jaringan WiFi lokal.'
    )
    doc.add_heading('Capaian Utama', level=2)
    capaian = [
        'Absensi dua faktor: RFID + Face Recognition untuk keamanan maksimal',
        'Pemrosesan real-time dengan notifikasi langsung melalui Socket.IO',
        'Registrasi otomatis untuk kartu RFID baru',
        'Laporan absensi yang dapat diekspor ke PDF',
        'Monitoring perangkat IoT secara real-time',
        'Antarmuka web modern dan responsif dengan React',
        'Dokumentasi lengkap untuk pengembangan dan pemeliharaan',
    ]
    for c in capaian:
        add_bullet(c)
    add_callout_box(
        'AttendTrack telah dirancang untuk kemudahan penggunaan, keandalan tinggi, '
        'dan skalabilitas — mendukung banyak perangkat IoT secara simultan.',
        'info', 'Tentang Sistem'
    )
    add_page_break()

# ── CHAPTER BUILDERS ──────────────────────────────────────────
def build_chapter_01():
    add_icon_heading('📘', 'Gambaran Umum Proyek')
    doc.add_paragraph(
        'AttendTrack adalah sistem manajemen absensi lengkap yang menggabungkan teknologi '
        'RFID (Radio Frequency Identification) dengan Kecerdasan Buatan Pengenalan Wajah '
        '(Face Recognition AI) untuk memberikan verifikasi absensi dua faktor. Sistem ini '
        'dirancang untuk institusi pendidikan, perkantoran, dan lingkungan manufaktur '
        'yang membutuhkan pencatatan absensi yang andal.'
    )
    add_callout_box(
        'Dengan kata lain, AttendTrack adalah sistem absensi pintar yang memastikan '
        'setiap kehadiran tercatat secara akurat karena membutuhkan dua bukti: '
        'kartu RFID DAN verifikasi wajah.',
        'info', 'Penjelasan Sederhana'
    )
    doc.add_heading('Fitur Utama', level=2)
    fitur = [
        'Absensi berbasis RFID: Memindai kartu RFID melalui ESP8266 + pembaca MFRC522',
        'Verifikasi Pengenalan Wajah: Autentikasi dua faktor dengan biometric wajah',
        'Monitoring Real-time: Event absensi langsung melalui Socket.IO',
        'Manajemen Karyawan: Operasi CRUD lengkap untuk data karyawan',
        'Manajemen Perangkat: Penemuan otomatis via mDNS, monitoring heartbeat',
        'Laporan Absensi: Ringkasan harian, pelacakan ketepatan waktu, ekspor PDF',
        'Korelasi Sesi: Menghubungkan event RFID + Wajah menggunakan jendela waktu',
        'Pendaftaran Otomatis: Kartu RFID baru memicu alur pengambilan wajah otomatis',
    ]
    for f in fitur:
        add_bullet(f)
    doc.add_heading('Komponen Hardware', level=2)
    add_table(
        ['Perangkat', 'Peran', 'Komponen Utama'],
        [
            ['ESP8266', 'Pemindai RFID Master', 'MFRC522 RFID, LCD 16x2 I2C, Buzzer'],
            ['ESP32-CAM', 'Kamera Pengambil Wajah', 'Kamera OV2640, LED Flash, PSRAM'],
            ['Server Backend', 'Server Aplikasi', 'Node.js + Python'],
        ],
        [1.5, 2.0, 2.5]
    )
    doc.add_heading('Tumpukan Teknologi', level=2)
    add_table(
        ['Lapisan', 'Teknologi', 'Tujuan'],
        [
            ['API Gateway', 'Express.js (Node.js) port 3000', 'REST API, Socket.IO, upload file'],
            ['gRPC Service', '@grpc/grpc-js port 50051', 'Logika bisnis internal'],
            ['Face Recognition', 'FastAPI (Python) port 8000', 'Enkoding wajah & verifikasi'],
        ],
        [2.0, 2.5, 1.5]
    )
    add_summary_box([
        'AttendTrack adalah sistem absensi dua faktor (RFID + Wajah)',
        'Terdiri dari 3 microservices: Gateway, gRPC, Face Recognition',
        'Didukung perangkat IoT: ESP8266 (RFID) dan ESP32-CAM (Kamera)',
        'Mendukung monitoring real-time dan laporan PDF',
    ])

def build_chapter_02():
    add_icon_heading('🏗', 'Arsitektur Sistem')
    add_horizontal_rule()
    doc.add_paragraph(
        'Sistem absensi menggunakan arsitektur microservices dengan tiga layanan '
        'utama dan dua komponen firmware IoT.'
    )
    add_callout_box(
        'Microservices adalah pendekatan di mana aplikasi dipecah menjadi '
        'layanan-layanan kecil yang independen. Setiap layanan memiliki tugas '
        'spesifik dan dapat dikembangkan serta dijalankan secara terpisah.',
        'info', 'Apa itu Microservices?'
    )
    doc.add_heading('Diagram Arsitektur', level=2)
    add_code('''
 LAPISAN KLIEN
 ┌──────────────────────┐   ┌──────────────────────────┐
 │   Web Frontend       │   │   Hardware IoT           │
 │   (React + Vite)     │   │   ESP8266 + ESP32-CAM    │
 │   Port 5173          │   │   (Jaringan Lokal)        │
 └────────┬─────────────┘   └────────┬─────────────────┘
          │ HTTP/WebSocket          │ HTTP (mDNS)
          ▼                          ▼
 ┌──────────────────────────────────────────────────────┐
 │              LAPISAN API GATEWAY                      │
 │              Express.js (Port 3000)                   │
 │  ┌─────────┐  ┌──────────┐  ┌──────────────────┐    │
 │  │ REST    │  │ Socket.IO│  │ Static/Uploads   │    │
 │  └────┬────┘  └──────────┘  └──────────────────┘    │
 └───────┼──────────────────────────────────────────────┘
         │ gRPC
         ▼
 ┌──────────────────────────────────────────────────────┐
 │           LAPISAN SERVICE (gRPC - Port 50051)         │
 │  Auth · Employee · Device · Attendance · Settings    │
 └───────┬──────────────────────────────────────────────┘
         │
 ┌───────┴──────────────────────────────────────────────┐
 │           LAPISAN DATA                                │
 │  ┌──────────────┐   ┌────────────────────────────┐   │
 │  │ SQLite       │   │ Python Face Recognition    │   │
 │  │ (SQL.js)     │   │ FastAPI (Port 8000)        │   │
 │  └──────────────┘   └────────────────────────────┘   │
 └──────────────────────────────────────────────────────┘''')
    doc.add_heading('Penjelasan Lapisan', level=2)
    doc.add_heading('Lapisan Klien', level=3)
    doc.add_paragraph(
        'Lapisan ini terdiri dari antarmuka pengguna berbasis web (React) '
        'dan perangkat IoT (ESP8266 dan ESP32-CAM).'
    )
    doc.add_heading('Lapisan API Gateway', level=3)
    doc.add_paragraph(
        'Bertindak sebagai pintu gerbang tunggal untuk semua permintaan. '
        'Gateway menangani routing API, koneksi WebSocket, dan penyajian file statis.'
    )
    add_summary_box([
        'Arsitektur microservices dengan 3 lapisan utama',
        'Komunikasi internal menggunakan gRPC',
        'Setiap layanan memiliki tanggung jawab yang terpisah',
        'Sistem mendukung banyak perangkat IoT secara simultan',
    ])

def build_chapter_03():
    add_icon_heading('⚙', 'Backend')
    doc.add_paragraph(
        'Backend adalah aplikasi TypeScript Node.js yang terstruktur dengan prinsip '
        'Clean Architecture. Terdiri dari dua komponen runtime yang berjalan dalam '
        'satu proses Node.js.'
    )
    add_callout_box(
        'Backend adalah "otak" dari sistem. Semua logika bisnis, validasi data, '
        'dan komunikasi antar komponen diproses di sini.',
        'info', 'Apa itu Backend?'
    )
    doc.add_heading('Komponen Runtime', level=2)
    add_table(
        ['Komponen', 'Port', 'Fungsi'],
        [
            ['Express Gateway', '3000', 'REST API + Socket.IO'],
            ['gRPC Server', '50051', 'Logika bisnis internal'],
        ],
        [2.0, 1.5, 2.5]
    )
    doc.add_heading('Struktur Proyek', level=2)
    add_code('''
src/
├── config/          # Konfigurasi environment
├── gateway/         # Express Gateway + Socket.IO
├── grpc/            # gRPC Server + Handlers
├── modules/         # Modul fitur (Clean Architecture)
│   ├── auth/        # Autentikasi
│   ├── employee/    # Manajemen karyawan
│   ├── device/      # Manajemen perangkat
│   ├── attendance/  # Absensi (modul terbesar)
│   └── settings/    # Pengaturan sistem
├── proto/           # Definisi protobuf
└── shared/          # Shared infrastructure''')
    add_callout_box(
        'Clean Architecture memisahkan kode menjadi lapisan-lapisan dengan '
        'tanggung jawab yang jelas, sehingga lebih mudah dipelihara dan dikembangkan.',
        'tip', 'Mengapa Clean Architecture?'
    )
    add_summary_box([
        'Backend menggunakan Node.js + TypeScript dengan Clean Architecture',
        'Express Gateway menangani REST API dan Socket.IO',
        'gRPC Server memproses logika bisnis',
        'Struktur modular memudahkan pengembangan fitur baru',
    ])

def build_chapter_04():
    add_icon_heading('🖥', 'Frontend')
    doc.add_paragraph(
        'Frontend adalah aplikasi React 18 single-page yang dibangun dengan '
        'Vite 5 dan TypeScript. Menyediakan dashboard modern dan responsif untuk '
        'mengelola sistem absensi.'
    )
    add_callout_box(
        'Frontend adalah "wajah" dari sistem — apa yang dilihat dan digunakan '
        'oleh pengguna setiap hari.',
        'info', 'Apa itu Frontend?'
    )
    doc.add_heading('Tumpukan Teknologi', level=2)
    add_table(
        ['Komponen', 'Teknologi', 'Versi'],
        [
            ['UI Framework', 'React + TypeScript', '18'],
            ['Build Tool', 'Vite', '5'],
            ['Styling', 'Tailwind CSS', '3'],
            ['State Management', 'Zustand', '5'],
            ['HTTP Client', 'Axios', '-'],
            ['Real-time', 'Socket.IO Client', '-'],
            ['Routing', 'React Router', 'v6'],
        ],
        [2.0, 2.5, 1.0]
    )
    doc.add_heading('Halaman Utama', level=2)
    halaman = [
        'Dashboard — Ringkasan absensi dan status perangkat',
        'History — Riwayat absensi dengan filter dan pencarian',
        'Sessions — Sesi absensi aktif',
        'Employees — Manajemen data karyawan',
        'Devices — Monitoring perangkat IoT',
        'Settings — Pengaturan sistem',
        'Reports — Laporan dan ekspor PDF',
    ]
    for h in halaman:
        add_bullet(h)
    add_summary_box([
        'Frontend menggunakan React 18 dengan Vite 5',
        'State management dengan Zustand',
        'Komunikasi real-time via Socket.IO',
        'Desain responsif dengan Tailwind CSS',
    ])

def build_chapter_05():
    add_icon_heading('🤖', 'Face Recognition')
    doc.add_paragraph(
        'Service Face Recognition adalah microservice Python FastAPI yang menyediakan '
        'kemampuan enkoding dan verifikasi wajah. Menggunakan DeepFace dengan model '
        'Facenet128 dan MediaPipe untuk deteksi wajah.'
    )
    add_callout_box(
        'Face Recognition adalah "mata" dari sistem — ia mengenali wajah '
        'setiap pengguna dengan membandingkannya dengan data yang sudah tersimpan.',
        'info', 'Cara Kerja Face Recognition'
    )
    doc.add_heading('Tumpukan Teknologi', level=2)
    add_table(
        ['Komponen', 'Teknologi'],
        [
            ['Web Framework', 'FastAPI (Python)'],
            ['Deteksi Wajah', 'MediaPipe'],
            ['Embedding Wajah', 'Facenet128 (via DeepFace)'],
            ['Pengolahan Citra', 'OpenCV'],
            ['Deep Learning', 'TensorFlow'],
        ],
        [2.5, 3.5]
    )
    doc.add_heading('Pipeline Face Recognition', level=2)
    add_code('''
 Input: Base64 Image
        │
        ▼
 1. Decode Base64 → OpenCV Image
        │
        ▼
 2. MediaPipe Face Detection
        │
        ▼
 3. Face Alignment & Crop
        │
        ▼
 4. Facenet128 Embedding
        │
        ▼
 5. Cosine Distance Calculation
        │
        ▼
 Output: { isMatch, confidence }''')
    add_summary_box([
        'Microservice Python untuk pengenalan wajah',
        'Menggunakan DeepFace + Facenet128 untuk akurasi tinggi',
        'MediaPipe untuk deteksi wajah',
        'Port 8000, berkomunikasi dengan backend via HTTP',
    ])

def build_chapter_06():
    add_icon_heading('📡', 'ESP8266 — Pemindai RFID')
    doc.add_paragraph(
        'Firmware ESP8266 berjalan di mikrokontroler ESP8266 dan berfungsi sebagai '
        'Pemindai RFID Master. Membaca kartu RFID, berkomunikasi dengan server backend, '
        'dan mengontrol ESP32-CAM melalui koneksi serial.'
    )
    doc.add_heading('Komponen Hardware', level=2)
    add_table(
        ['Komponen', 'Spesifikasi'],
        [
            ['MCU', 'ESP8266'],
            ['RFID Reader', 'MFRC522 (SPI)'],
            ['Display', 'LCD 16x2 I2C (address 0x27)'],
            ['Buzzer', 'Active buzzer GPIO0'],
            ['Serial', 'SoftwareSerial untuk ESP32-CAM'],
        ],
        [2.0, 4.0]
    )
    doc.add_heading('Pemetaan Pin', level=2)
    add_table(
        ['Komponen', 'Pin ESP8266', 'GPIO'],
        [
            ['MFRC522 SDA', 'D8', 'GPIO15'],
            ['MFRC522 RST', 'D4', 'GPIO2'],
            ['MFRC522 MOSI', 'D7', 'GPIO13'],
            ['MFRC522 MISO', 'D6', 'GPIO12'],
            ['MFRC522 SCK', 'D5', 'GPIO14'],
            ['LCD I2C SDA', 'D2', 'GPIO4'],
            ['LCD I2C SCL', 'D1', 'GPIO5'],
            ['Buzzer', 'D3', 'GPIO0'],
        ],
        [2.5, 2.0, 2.0]
    )
    add_summary_box([
        'ESP8266 bertugas membaca kartu RFID',
        'Berkomunikasi dengan backend via HTTP',
        'Mengontrol ESP32-CAM via koneksi serial',
        'Dilengkapi LCD untuk tampilan status',
    ])

def build_chapter_07():
    add_icon_heading('📷', 'ESP32-CAM — Kamera Wajah')
    doc.add_paragraph(
        'Firmware ESP32-CAM berjalan di modul ESP32-CAM dan berfungsi sebagai '
        'Kamera Pengambil Wajah. Menerima perintah dari ESP8266 melalui koneksi '
        'serial, mengambil gambar wajah, dan mengirimnya ke server backend via HTTP.'
    )
    doc.add_heading('Komponen Hardware', level=2)
    add_table(
        ['Komponen', 'Spesifikasi'],
        [
            ['MCU', 'ESP32 (ESP32-CAM module)'],
            ['Kamera', 'OV2640 (2MP)'],
            ['LED Flash', 'GPIO4'],
            ['PSRAM', 'Opsional (2MB)'],
            ['Serial', 'Hardware Serial2'],
        ],
        [2.0, 4.0]
    )
    add_callout_box(
        'ESP32-CAM bekerja sama dengan ESP8266: ESP8266 membaca kartu RFID, '
        'lalu memberi perintah ke ESP32-CAM untuk mengambil foto wajah.',
        'tip', 'Kerja Sama ESP8266 dan ESP32-CAM'
    )
    add_summary_box([
        'ESP32-CAM bertugas mengambil foto wajah',
        'Dilengkapi LED flash untuk pencahayaan',
        'Berkomunikasi dengan ESP8266 via serial',
        'Mengirim gambar ke backend via HTTP',
    ])

def build_chapter_08():
    add_icon_heading('🗄', 'Database')
    doc.add_paragraph(
        'Sistem menggunakan SQL.js — implementasi JavaScript dari SQLite yang '
        'dikompilasi ke WebAssembly. Ini menghilangkan kebutuhan akan server database '
        'eksternal.'
    )
    add_callout_box(
        'Database adalah "memori" dari sistem — semua data karyawan, '
        'absensi, dan pengaturan disimpan di sini.',
        'info', 'Apa itu Database?'
    )
    doc.add_heading('Karakteristik Utama', level=2)
    karakter = [
        'File database tunggal: storage/rfid_v3.sqlite',
        'Eksekusi in-memory dengan persistensi file otomatis',
        'Mode WAL untuk performa baca konkuren',
        'Tidak ada dependensi eksternal',
        'Otomatis dibuat pada first run',
    ]
    for k in karakter:
        add_bullet(k)
    doc.add_heading('Entity Relationship Diagram', level=2)
    add_code('''
 Tabel-tabel dalam sistem:
 ┌──────────────┐    ┌──────────────────┐    ┌──────────────┐
 │    USERS     │    │   EMPLOYEES      │    │   DEVICES    │
 │──────────────│    │──────────────────│    │──────────────│
 │ id           │    │ id               │    │ id           │
 │ name         │    │ full_name        │    │ device_code  │
 │ email        │    │ department       │    │ type         │
 │ password_hash│    │ position         │    │ name         │
 │ role         │    │ rfid_uid         │    │ location     │
 └──────────────┘    │ face_descriptor  │    │ status       │
                     │ face_image_path  │    │ last_seen_at │
                     │ is_active        │    └──────────────┘
                     └──────────────────┘

 ┌──────────────────────────┐    ┌──────────────────────────┐
 │   ATTENDANCE_SESSIONS    │    │   ATTENDANCE_RECORDS     │
 │──────────────────────────│    │──────────────────────────│
 │ id                       │    │ id                       │
 │ correlation_id           │    │ session_id ──────────┐   │
 │ pairing_key              │    │ employee_id ───────┐ │   │
 │ rfid_uid                 │    │ rfid_uid           │ │   │
 │ status                   │    │ status             │ │   │
 │ started_at               │    │ confidence         │ │   │
 └──────────────────────────┘    │ image_path         │ │   │
                                 │ verified_at        │ │   │
                                 └────────────────────┘─┘───┘''')
    add_summary_box([
        'Menggunakan SQL.js (SQLite in-browser)',
        'Tanpa server database eksternal',
        '6 tabel utama: Users, Employees, Devices, Sessions, Records, Settings',
        'Otomatis dibuat dan dimigrasi pada first run',
    ])

def build_chapter_09():
    add_icon_heading('🔄', 'API')
    doc.add_paragraph(
        'Semua endpoint API dilayani oleh Express Gateway pada port 3000. '
        'Endpoint diawali dengan /api/v1/.'
    )
    add_callout_box(
        'API (Application Programming Interface) adalah "jembatan" yang '
        'memungkinkan aplikasi-aplikasi berbeda untuk saling berkomunikasi.',
        'info', 'Apa itu API?'
    )
    doc.add_heading('Endpoint Autentikasi', level=2)
    add_table(
        ['Metode', 'Endpoint', 'Deskripsi', 'Auth'],
        [
            ['POST', '/auth/login', 'Login user', 'Tidak'],
            ['POST', '/auth/register', 'Registrasi user', 'Tidak'],
            ['GET', '/auth/me', 'Profil user', 'JWT'],
        ],
        [1.0, 1.5, 2.0, 0.8]
    )
    doc.add_heading('Endpoint Karyawan', level=2)
    add_table(
        ['Metode', 'Endpoint', 'Deskripsi', 'Auth'],
        [
            ['GET', '/employees', 'Daftar karyawan', 'JWT'],
            ['POST', '/employees', 'Tambah karyawan', 'JWT'],
            ['PUT', '/employees/:id', 'Update karyawan', 'JWT'],
            ['DELETE', '/employees/:id', 'Hapus karyawan', 'JWT'],
        ],
        [1.0, 1.8, 2.2, 0.8]
    )
    doc.add_heading('Endpoint Absensi', level=2)
    add_table(
        ['Metode', 'Endpoint', 'Deskripsi', 'Auth'],
        [
            ['POST', '/attendance/rfid', 'Proses RFID', 'Tidak'],
            ['POST', '/attendance/face', 'Proses wajah', 'Tidak'],
            ['GET', '/attendance/history', 'Riwayat absensi', 'JWT'],
            ['GET', '/attendance/sessions', 'Sesi aktif', 'JWT'],
            ['GET', '/attendance/stream', 'Stream real-time', 'JWT'],
        ],
        [1.0, 2.0, 2.0, 0.8]
    )
    add_summary_box([
        'REST API dengan prefix /api/v1/',
        'Express Gateway di port 3000',
        'Autentikasi menggunakan JWT Bearer token',
        'Endpoint publik untuk perangkat IoT',
    ])

def build_chapter_10():
    add_icon_heading('⚡', 'Socket.IO — Komunikasi Real-time')
    doc.add_paragraph(
        'Socket.IO menyediakan komunikasi real-time bidirectional antara server backend '
        'dan aplikasi frontend web. Digunakan untuk update absensi langsung, perubahan '
        'status perangkat, dan event registrasi.'
    )
    doc.add_heading('Arsitektur', level=2)
    add_code('''
 Frontend React ◄──── Socket.IO ────► Express Gateway ◄──── EventEmitter
                           │
                    attendance:new
                    attendance:update
                    rfid:new
                    registration:image
                    device:status
                    session:created''')
    doc.add_heading('Event Flow', level=2)
    doc.add_paragraph(
        'Keitka terjadi event absensi, alurnya adalah sebagai berikut:'
    )
    add_code('''
 1. ESP8266 POST RFID ──► Gateway
 2. Gateway buat sesi ──► emit session:created
 3. ESP32-CAM POST face ──► Gateway
 4. Gateway tambah face ke sesi ──► emit attendance:update
 5. Verifikasi selesai ──► emit attendance:new
 6. Frontend terima ──► Update LiveFeed + Toast notification''')
    add_summary_box([
        'Komunikasi real-time dua arah',
        'WebSocket dengan fallback HTTP long-polling',
        '6 event utama untuk update absensi dan perangkat',
        'Integrasi dengan EventEmitter internal backend',
    ])

def build_chapter_11():
    add_icon_heading('🔗', 'gRPC')
    doc.add_paragraph(
        'gRPC digunakan untuk komunikasi internal service-to-service antara Express '
        'Gateway dan lapisan logika bisnis. Menyediakan kontrak yang strongly-typed '
        'melalui Protocol Buffers.'
    )
    add_callout_box(
        'gRPC adalah sistem komunikasi berkinerja tinggi yang memungkinkan '
        'layanan-layanan berbicara satu sama lain dengan cepat dan efisien.',
        'info', 'Apa itu gRPC?'
    )
    doc.add_heading('Service Definitions', level=2)
    add_table(
        ['Service', 'Fungsi Utama'],
        [
            ['AuthService', 'Login, Register, Validate Token'],
            ['EmployeeService', 'CRUD Karyawan, FindByRFID'],
            ['DeviceService', 'Register, Heartbeat, FindAll'],
            ['AttendanceService', 'CheckRFID, HandleFace, History, Sessions'],
            ['SettingsService', 'Get, Upsert, Reset Settings'],
        ],
        [2.5, 3.5]
    )
    add_summary_box([
        'Komunikasi internal menggunakan gRPC',
        'Port 50051 dengan protokol HTTP/2',
        '5 service definitions dalam platform.proto',
        'Kontak kuat dengan TypeScript types',
    ])

def build_chapter_12():
    add_icon_heading('📋', 'Alur Aplikasi')
    doc.add_paragraph(
        'Bab ini menjelaskan alur lengkap pengguna dalam menggunakan sistem, '
        'dari login hingga absensi.'
    )
    doc.add_heading('Alur Login', level=2)
    add_code('''
 1. User buka browser → http://{server}:5173
 2. Cek token di localStorage
 3. Jika tidak ada → redirect ke /login
 4. User masukkan email & password
 5. Frontend POST /api/v1/auth/login
 6. Validasi kredensial (bcrypt)
 7. Generate JWT token
 8. Simpan token → redirect ke Dashboard
 9. Load data awal: employees, devices, sessions, history
10. Connect Socket.IO untuk update real-time''')
    doc.add_heading('Alur Absensi RFID', level=2)
    add_code('''
 1. User tap kartu RFID di ESP8266
 2. ESP8266 baca UID → POST ke backend
 3. Backend cek RFID:
    - TERDAFTAR → Minta capture wajah
    - TIDAK TERDAFTAR → Minta registrasi wajah
 4. ESP32-CAM ambil foto wajah
 5. Backend verifikasi wajah (DeepFace)
 6. Simpan record absensi
 7. Kirim notifikasi real-time ke Frontend
 8. ESP8266 tampilkan hasil di LCD''')
    add_summary_box([
        'Login memerlukan email + password dengan JWT',
        'Absensi melalui RFID + Face Recognition',
        'Proses real-time dengan notifikasi langsung',
        'Registrasi otomatis untuk kartu baru',
    ])

def build_chapter_13():
    add_icon_heading('📊', 'Alur Fitur')
    doc.add_heading('Fitur Login', level=2)
    doc.add_paragraph(
        'Fitur ini memungkinkan pengguna yang berwenang untuk mengakses sistem. '
        'Proses autentikasi dilakukan dengan email dan password.'
    )
    doc.add_heading('Komponen Terlibat:', level=3)
    add_bullet('Frontend: login-page.tsx, auth-store.ts, auth.service.ts')
    add_bullet('Backend: auth.controller.ts, auth.service.ts, user.repository.ts')
    add_bullet('gRPC: auth.handler.ts, AuthService.Login')
    doc.add_heading('Fitur Dashboard', level=2)
    doc.add_paragraph(
        'Dashboard adalah halaman utama yang menampilkan ringkasan absensi '
        'dan status sistem secara real-time. Informasi yang ditampilkan meliputi '
        'jumlah kehadiran valid, scan tidak valid, total scan, dan status perangkat IoT.'
    )
    doc.add_heading('Fitur Laporan', level=2)
    doc.add_paragraph(
        'Sistem menyediakan laporan absensi yang dapat difilter berdasarkan '
        'tanggal dan departemen. Laporan dapat diekspor ke format PDF.'
    )
    add_summary_box([
        'Login dengan autentikasi JWT',
        'Dashboard menampilkan data real-time',
        'Laporan dapat diekspor ke PDF',
        'Manajemen karyawan dan perangkat',
    ])

def build_chapter_14():
    add_icon_heading('🔀', 'Diagram Sekuens')
    doc.add_paragraph(
        'Diagram sekuens berikut menjelaskan interaksi antar komponen '
        'dalam setiap skenario penggunaan sistem.'
    )
    doc.add_heading('Login', level=2)
    add_code('''
 User → Frontend: Buka aplikasi
 Frontend → Frontend: Cek token localStorage
 alt Token tidak ada
     Frontend → User: Redirect ke /login
     User → Frontend: Isi email & password
     Frontend → Gateway: POST /api/v1/auth/login
     Gateway → gRPC: AuthService.Login
     gRPC → DB: SELECT user by email
     alt Valid
         gRPC → gRPC: Generate JWT
         gRPC → Gateway: { token, user }
         Gateway → Frontend: { token, user }
         Frontend → Frontend: Simpan token
     else Invalid
         Gateway → Frontend: Error 401
 else Token ada
     Frontend → Gateway: GET /auth/me
     alt Token valid → Dashboard
     else Expired → Redirect /login''')
    doc.add_heading('Absensi RFID (Kartu Terdaftar)', level=2)
    add_code('''
 User → ESP8266: Tap kartu RFID
 ESP8266 → Gateway: POST /check-rfid
 Gateway → gRPC: AttendanceService.CheckRfid
 gRPC → DB: Cari karyawan by rfid_uid
 Gateway → ESP8266: { registered: true }
 ESP8266 → ESP32-CAM: CAPTURE command
 ESP32-CAM → Gateway: POST /face (dengan image)
 Gateway → gRPC: AttendanceService.HandleFace
 gRPC → Python: POST /verify
 Python → Python: Encode + compare
 Python → gRPC: { isMatch: true, confidence }
 gRPC → DB: INSERT attendance record
 gRPC → Socket: Emit attendance:new
 Socket → Frontend: Update LiveFeed + Summary
 ESP8266 → LCD: "Absensi tercatat"''')
    add_summary_box([
        'Diagram sekuens untuk Login, Absensi, Registrasi',
        'Menunjukkan interaksi antar semua komponen',
        'Setiap skenario memiliki alur yang jelas',
        'Memudahkan pemahaman sistem secara visual',
    ])

def build_chapter_15():
    add_icon_heading('🚀', 'Deployment')
    doc.add_paragraph(
        'Panduan ini menjelaskan langkah-langkah untuk menerapkan sistem '
        'di lingkungan production.'
    )
    doc.add_heading('Persyaratan Sistem', level=2)
    add_table(
        ['Komponen', 'Persyaratan'],
        [
            ['OS Server', 'Windows 10+, Linux, macOS'],
            ['Node.js', 'v18.x atau lebih baru'],
            ['Python', '3.10 atau lebih baru'],
            ['RAM', 'Minimal 4GB (8GB recommended)'],
            ['Storage', 'Minimal 1GB free'],
        ],
        [2.0, 4.0]
    )
    doc.add_heading('Langkah Deployment', level=2)
    steps = [
        'Clone repository dan install dependensi',
        'Konfigurasi environment variables',
        'Instal dan jalankan Face Recognition Service',
        'Jalankan Backend (npm run dev)',
        'Deploy Frontend (npm run build)',
        'Upload firmware ke ESP8266 dan ESP32-CAM',
    ]
    for s in steps:
        add_bullet(s)
    add_callout_box(
        'Untuk deployment production, gunakan static IP untuk server backend, '
        'konfigurasi firewall yang tepat, dan backup database secara berkala.',
        'warning', 'Catatan Deployment'
    )
    add_summary_box([
        'Deployment backend: Node.js + Python',
        'Frontend di-build dan di-deploy ke web server',
        'Firmware ESP8266 dikonfigurasi via WiFi portal',
        'Semua komponen berjalan di jaringan lokal yang sama',
    ])

def build_chapter_16():
    add_icon_heading('🔧', 'Troubleshooting')
    doc.add_paragraph(
        'Panduan ini membantu mengatasi masalah umum yang mungkin '
        'terjadi pada sistem.'
    )
    doc.add_heading('Masalah Backend', level=2)
    doc.add_heading('Server Tidak Mau Start', level=3)
    add_table(
        ['Gejala', 'Penyebab', 'Solusi'],
        [
            ['EADDRINUSE', 'Port sudah digunakan', 'Matikan proses di port tsb'],
            ['JWT_SECRET required', 'Env tidak lengkap', 'Lengkapi file .env'],
            ['SQLITE_BUSY', 'Database terkunci', 'Restart server'],
        ],
        [2.0, 1.5, 2.5]
    )
    doc.add_heading('Masalah ESP8266', level=2)
    add_table(
        ['Gejala', 'Penyebab', 'Solusi'],
        [
            ['Tidak connect WiFi', 'SSID/Password salah', 'Konfigurasi ulang via portal'],
            ['Tidak baca RFID', 'Koneksi longgar', 'Periksa kabel SPI'],
            ['LCD tidak tampil', 'Alamat I2C salah', 'Coba 0x27 atau 0x3F'],
        ],
        [2.0, 1.8, 2.5]
    )
    doc.add_heading('Masalah Face Recognition', level=2)
    add_table(
        ['Gejala', 'Penyebab', 'Solusi'],
        [
            ['False Positive', 'Threshold terlalu rendah', 'Naikkan threshold'],
            ['False Negative', 'Threshold terlalu tinggi', 'Turunkan threshold'],
            ['Lambat', 'CPU 100%', 'Gunakan GPU atau turunkan resolusi'],
        ],
        [2.0, 2.0, 2.5]
    )
    add_summary_box([
        'Backend: cek port, env variables, database',
        'ESP8266: periksa WiFi, kabel, alamat I2C',
        'ESP32-CAM: periksa GPIO0 saat flashing',
        'Face Recognition: atur threshold sesuai kebutuhan',
    ])

def build_chapter_17():
    add_icon_heading('🏆', 'Praktik Terbaik')
    doc.add_heading('Keamanan', level=2)
    keamanan = [
        'Ganti JWT_SECRET default dengan nilai acak',
        'Jangan commit file .env ke repository',
        'Validasi semua input menggunakan Zod schema',
        'Gunakan pairing key yang unik untuk setiap device',
        'Aktifkan Helmet middleware untuk security headers',
    ]
    for k in keamanan:
        add_bullet(k)
    doc.add_heading('Performa', level=2)
    performa = [
        'Pre-load model Face Recognition saat startup',
        'Downscale gambar sebelum dikirim ke face service',
        'Gunakan lazy loading untuk komponen frontend',
        'Batch heartbeat interval 30-60 detik',
    ]
    for p in performa:
        add_bullet(p)
    add_callout_box(
        'Selalu backup database secara berkala dan uji '
        'proses recovery untuk memastikan data aman.',
        'warning', 'Penting'
    )
    add_summary_box([
        'Keamanan: JWT, env vars, input validation',
        'Performa: caching, lazy loading, batch processing',
        'Maintenance: backup rutin, update firmware',
        'Reliabilitas: watchdog timer, retry logic',
    ])

# ── KESIMPULAN ───────────────────────────────────────────────
def build_kesimpulan():
    doc.add_heading('Kesimpulan', level=1)
    doc.add_paragraph(
        'AttendTrack — Sistem Absensi RFID + Face Recognition versi 3.0 telah berhasil '
        'dikembangkan dengan seluruh fitur utama berjalan dengan baik. Sistem ini '
        'menggabungkan teknologi RFID dan kecerdasan buatan untuk memberikan solusi '
        'absensi dua faktor yang andal, akurat, dan real-time.'
    )
    doc.add_heading('Keunggulan Sistem', level=2)
    keunggulan = [
        'Akurasi tinggi dengan verifikasi dua faktor (RFID + Wajah)',
        'Pemrosesan real-time dengan notifikasi langsung',
        'Registrasi otomatis untuk kartu baru',
        'Antarmuka yang mudah digunakan',
        'Dokumentasi lengkap untuk pengembangan',
        'Arsitektur modular yang mudah dikembangkan',
    ]
    for k in keunggulan:
        add_bullet(k)

# ── LAMPIRAN ──────────────────────────────────────────────────
def build_lampiran():
    doc.add_heading('Lampiran', level=1)
    doc.add_heading('A. Glosarium', level=2)
    add_table(
        ['Istilah', 'Artinya'],
        [
            ['RFID', 'Radio Frequency Identification — teknologi identifikasi via gelombang radio'],
            ['gRPC', 'Remote Procedure Call — protokol komunikasi internal'],
            ['JWT', 'JSON Web Token — token autentikasi'],
            ['mDNS', 'Multicast DNS — penemuan perangkat dalam jaringan'],
            ['WebSocket', 'Protokol komunikasi real-time dua arah'],
            ['API', 'Application Programming Interface — antarmuka pemrograman'],
            ['Microservices', 'Arsitektur dengan layanan-layanan kecil yang independen'],
        ],
        [2.0, 4.5]
    )
    doc.add_heading('B. Referensi', level=2)
    refs = [
        'Dokumentasi Express.js — https://expressjs.com',
        'Dokumentasi React — https://react.dev',
        'Dokumentasi Socket.IO — https://socket.io',
        'Dokumentasi gRPC — https://grpc.io',
        'Dokumentasi FastAPI — https://fastapi.tiangolo.com',
        'Dokumentasi DeepFace — https://github.com/serengil/deepface',
    ]
    for r in refs:
        add_bullet(r)

# ── BUILD DOCUMENT ────────────────────────────────────────────
def main():
    build_cover()
    build_kata_pengantar()
    build_toc()
    build_executive_summary()

    chapters = [
        build_chapter_01,
        build_chapter_02,
        build_chapter_03,
        build_chapter_04,
        build_chapter_05,
        build_chapter_06,
        build_chapter_07,
        build_chapter_08,
        build_chapter_09,
        build_chapter_10,
        build_chapter_11,
        build_chapter_12,
        build_chapter_13,
        build_chapter_14,
        build_chapter_15,
        build_chapter_16,
        build_chapter_17,
    ]
    for i, chapter_func in enumerate(chapters, 1):
        add_page_break()
        chapter_func()

    add_page_break()
    build_kesimpulan()
    add_page_break()
    build_lampiran()

    # ── Headers & Footers ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run('AttendTrack — Dokumentasi Sistem v3.0')
        run.font.size = Pt(8)
        run.font.color.rgb = GREY
        run.font.name = 'Calibri'
        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('Halaman ')
        run.font.size = Pt(8)
        run.font.color.rgb = GREY
        # Add page number field
        fldChar1 = parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        run2 = fp.add_run()
        run2._r.append(fldChar1)
        instrText = parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        run3 = fp.add_run()
        run3._r.append(instrText)
        fldChar2 = parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run4 = fp.add_run()
        run4._r.append(fldChar2)

    doc.save(OUTPUT_FILE)
    print(f"Dokumen berhasil dibuat: {OUTPUT_FILE}")

if __name__ == '__main__':
    main()
